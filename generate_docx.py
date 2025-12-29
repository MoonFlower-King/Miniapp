from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()
    
    # Title
    title = document.add_heading('知岁 (ZhiSui) - 产品实验报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    p = document.add_paragraph()
    p.add_run('项目名称: ').bold = True
    p.add_run('知岁 (ZhiSui) - 原 PocketLedger\n')
    p.add_run('开发者: ').bold = True
    p.add_run('[您的姓名]\n')
    p.add_run('开发日期: ').bold = True
    p.add_run('2025年12月\n')
    p.add_run('版本: ').bold = True
    p.add_run('v1.0.0')

    document.add_heading('1. 产品功能介绍', level=1)
    
    document.add_paragraph('“知岁”取“知晓岁月”之意，是一款集“个人记账”与“高效任务管理”于一体的 Android 效率工具。它通过极简的交互和强大的 AI 辅助，帮助用户更好地规划生活与财务。')
    
    document.add_heading('1.1 核心功能模块', level=2)
    
    # Ledger
    document.add_heading('💰 知岁账本 (Smart Ledger)', level=3)
    p = document.add_paragraph(style='List Bullet')
    p.add_run('多维记录: ').bold = True
    p.add_run('支持“支出”与“收入”双向记录，内置餐饮、交通、购物等 10+ 种常用分类。')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('可视化报表: ').bold = True
    p.add_run('首页动态展示收支圆环图 (Ring Chart)，配合百分比统计，财务状况一目了然。')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('流水明细: ').bold = True
    p.add_run('清晰的时间轴流水列表，支持长按删除与编辑。')

    # Tasks
    document.add_heading('✅ 知岁清单 (Pro Tasks)', level=3)
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Notion 风格管理: ').bold = True
    p.add_run('每一条任务都拥有丰富的属性，包括状态、优先级、截止日期、标签、负责人及附件。')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('可视性控制: ').bold = True
    p.add_run('独创“属性可见性”开关，用户可自定义列表展示哪些字段，保持界面清爽。')

    # AI
    document.add_heading('🤖 AI 智能助手 (AI Assistant)', level=3)
    p = document.add_paragraph(style='List Bullet')
    p.add_run('自然语言记账: ').bold = True
    p.add_run('输入“吃面花了25元”，AI 自动生成账单。')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('智能任务创建: ').bold = True
    p.add_run('输入“周五交报告”，AI 自动创建带截止日期的任务。')

    document.add_heading('2. 程序概要设计', level=1)
    
    document.add_heading('2.1 交互设计', level=2)
    document.add_paragraph('应用采用 Bottom Navigation + ViewPager2 的主流架构，实现了左右滑动切换三大核心模块（账本、任务、AI）丝滑体验。')
    
    document.add_heading('2.2 数据存储设计', level=2)
    document.add_paragraph('使用 Android 原生 SQLite 数据库进行本地离线存储。')
    document.add_paragraph('主要表结构：')
    document.add_paragraph('1. transactions (账单表): id, amount, type, category, note, date')
    document.add_paragraph('2. todos_v2 (任务表): id, title, status, priority, due_date, assignee, attachment_path')

    document.add_heading('3. 软件架构图', level=1)
    document.add_paragraph('（此处为 Mermaid 架构图的文字描述，详细图表请见附件或 Markdown 版本）')
    document.add_paragraph('架构模式：MVC (Model-View-Controller) + Event-Driven')
    document.add_paragraph('UI层: Activity/Fragment -> Logic层: Adapter/Manager -> Data层: SQLite/API')

    document.add_heading('4. 技术亮点与实现原理', level=1)
    
    document.add_heading('4.1 智能语义解析', level=2)
    document.add_paragraph('利用 OkHttp 对接 DeepSeek V3 大模型 API。配合 Prompt Engineering 强制 AI 输出标准 JSON 格式数据，再经由 Gson 解析写入数据库。')
    
    document.add_heading('4.2 工业级列表优化', level=2)
    document.add_paragraph('全面采用 RecyclerView + ViewHolder 复用机制。并在 Adapter 中实现了动态 View 显隐逻辑，以支持自定义属性可见性。')
    
    document.add_heading('4.3 数据库无缝迁移', level=2)
    document.add_paragraph('利用 SQLiteOpenHelper 的 onUpgrade 机制，编写 SQL 脚本 (ALTER TABLE) 实现从 v7 到 v8 的无损字段扩展。')

    document.save('product_report.docx')
    print("Word document generated successfully: product_report.docx")

if __name__ == "__main__":
    create_report()
